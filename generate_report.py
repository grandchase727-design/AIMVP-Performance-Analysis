"""AIMVP 정정 요약 리포트 생성기 (generate_report.py)

연속체인 정정 엔진(`simulate_portfolio`, 2026-06-23 정정) 기준으로 핵심 수치·차트·
요약 docx를 **재생성**한다. 기존 정적 리포트(docx)와 차트 생성 스크립트가 repo에
존재하지 않아(백로그의 "생성기 부재") 이 파일이 그 역할을 대체한다.

사용:
    py generate_report.py                 # Report Date = 오늘(전체 기간)
    py generate_report.py --date 2026-05-22

산출물(작업 폴더):
    report_적극형_cumulative.png
    report_중립형_cumulative.png
    report_swap_sweep.png
    AIMVP_정정요약_<date>.docx

※ docx/PNG 출력물은 .gitignore에 의해 추적되지 않음(재생성 가능). 본 스크립트만 커밋 대상.
"""
import sys
import os
import types
import argparse
from datetime import datetime

import matplotlib as mpl
mpl.use('Agg')
import matplotlib.pyplot as plt
# 한글 폰트 (Windows: 맑은 고딕). 없으면 기본 폰트로 폴백.
for _f in ['Malgun Gothic', 'AppleGothic', 'NanumGothic']:
    try:
        mpl.font_manager.findfont(_f, fallback_to_default=False)
        mpl.rcParams['font.family'] = _f
        break
    except Exception:
        continue
mpl.rcParams['axes.unicode_minus'] = False

import pandas as pd
import numpy as np

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

APP_DIR = os.path.dirname(os.path.abspath(__file__))
APP_PATH = os.path.join(APP_DIR, 'aimvp_streamlit_app.py')

# ── 색상 (대시보드 표준과 동일) ──
C_AGG, C_NEU = '#1F3A68', '#C48D43'
C_SWAP, C_ACWI = '#DC2626', '#9CA3AF'
HEAD1, HEAD2 = RGBColor(0x1F, 0x3A, 0x68), RGBColor(0x2E, 0x5A, 0x88)


def load_app(app_path=APP_PATH):
    """Streamlit 실행 없이 app 모듈의 함수/상수를 로드 (Sidebar 직전까지 exec)."""
    st = types.ModuleType('streamlit')

    def _noop(*a, **k):
        if len(a) == 1 and callable(a[0]) and not k:
            return a[0]
        def deco(fn=None):
            return (lambda f: f) if fn is None else fn
        return deco

    class _Any:
        def __getattr__(self, n):
            return _noop
        def __call__(self, *a, **k):
            return _noop(*a, **k)

    for n in ['set_page_config', 'markdown', 'caption', 'cache_data', 'cache_resource']:
        setattr(st, n, _Any())
    st.session_state = {}
    sys.modules['streamlit'] = st

    src = open(app_path, encoding='utf-8').read()
    head = src[:src.index('# ============== Sidebar ==============')]
    g = {'__file__': app_path, '__name__': 'aimvp_app'}
    exec(compile(head, app_path, 'exec'), g)
    return g


def compute_all(g, end_date_str):
    fetch_prices = g['fetch_prices']
    parse_holdings = g['parse_holdings']
    simulate_portfolio = g['simulate_portfolio']
    compute_metrics = g['compute_metrics']
    compute_bull_hit_rate = g['compute_bull_hit_rate']

    end_ts = pd.Timestamp(end_date_str)
    px = fetch_prices(end_date_str)
    out = {'px': px, 'end_ts': end_ts, 'price_last': px.index.max()}
    for key, label in [('적극', '적극형'), ('중립', '중립형')]:
        reb = [r for r in parse_holdings(key) if pd.Timestamp(r[0]) <= end_ts]
        actual = simulate_portfolio(reb, px, 0.0, end_date=end_date_str)
        swap50 = simulate_portfolio(reb, px, 0.5, end_date=end_date_str)
        ma, ms = compute_metrics(actual), compute_metrics(swap50)
        hits, n_bull, avg_excess, _ = compute_bull_hit_rate(reb, px)
        out[label] = {
            'rebals': reb, 'actual': actual, 'swap50': swap50, 'm_actual': ma, 'm_swap': ms,
            'itd_actual': ma['ITD'], 'itd_swap': ms['ITD'], 'alpha': ms['ITD'] - ma['ITD'],
            'sharpe_actual': ma['Sharpe'], 'sharpe_swap': ms['Sharpe'],
            'cagr_actual': ma['CAGR'], 'mdd_actual': ma['MaxDD'], 'vol_actual': ma['Vol'],
            'hits': hits, 'n_bull': n_bull, 'avg_excess': avg_excess,
        }
    return out


def _cum_from_daily(daily, idx):
    s = daily.reindex(idx).fillna(0.0)
    return ((1 + s).cumprod() - 1) * 100


def chart_cumulative(data, profile, outpath):
    px = data['px']
    d = data[profile]
    reb = d['rebals']
    start = pd.Timestamp(reb[0][0])
    end_ts = data['end_ts']
    idx = px.index[(px.index >= start) & (px.index <= end_ts)]
    ca = _cum_from_daily(d['actual'], idx)
    cs = _cum_from_daily(d['swap50'], idx)
    acwi = ((1 + px['ACWI'].reindex(idx).pct_change().fillna(0)).cumprod() - 1) * 100
    color = C_AGG if profile == '적극형' else C_NEU

    fig, ax = plt.subplots(figsize=(10, 5))
    ax.plot(idx, acwi.values, label='ACWI 100% (참고)', color=C_ACWI, lw=1.5, ls='--')
    ax.plot(idx, ca.values, label=f'실제 (스왑 0%) — 종점 {ca.iloc[-1]:+.1f}%', color=color, lw=2.4)
    ax.plot(idx, cs.values, label=f'50% 스왑 룰 — 종점 {cs.iloc[-1]:+.1f}%', color=C_SWAP, lw=2.0)
    ax.axhline(0, color='gray', lw=0.8, ls=':')
    ax.set_title(f'{profile} 누적 성과 (연속체인 정정 엔진)  ·  {idx[0].date()} ~ {idx[-1].date()}',
                 fontsize=13, color='#1F3A68')
    ax.set_ylabel('누적 수익률 (%)')
    ax.legend(loc='upper left', fontsize=10)
    ax.grid(alpha=0.3)
    fig.tight_layout()
    fig.savefig(outpath, dpi=130)
    plt.close(fig)


def chart_swap_sweep(data, outpath):
    g_sim = data['_g']['simulate_portfolio']
    g_met = data['_g']['compute_metrics']
    px = data['px']
    end_str = data['end_str']
    sw = [i / 10 for i in range(11)]
    x = [int(s * 100) for s in sw]
    res = {}
    for profile in ['적극형', '중립형']:
        reb = data[profile]['rebals']
        itd, sh = [], []
        for s in sw:
            m = g_met(g_sim(reb, px, s, end_date=end_str))
            itd.append(m['ITD'])
            sh.append(m['Sharpe'])
        res[profile] = {'itd': itd, 'sh': sh}

    fig, axes = plt.subplots(1, 2, figsize=(12, 4.6))
    for profile, c in [('적극형', C_AGG), ('중립형', C_NEU)]:
        axes[0].plot(x, res[profile]['itd'], marker='o', ms=4, label=profile, color=c)
        axes[1].plot(x, res[profile]['sh'], marker='o', ms=4, label=profile, color=c)
    axes[0].set_title('누적 수익률(ITD) vs 스왑 비율', fontsize=12, color='#1F3A68')
    axes[0].set_xlabel('SPY→QQQ 스왑 비율 (%)')
    axes[0].set_ylabel('ITD 누적 (%)')
    axes[1].set_title('Sharpe vs 스왑 비율', fontsize=12, color='#1F3A68')
    axes[1].set_xlabel('SPY→QQQ 스왑 비율 (%)')
    axes[1].set_ylabel('Sharpe')
    for ax in axes:
        ax.axvline(50, color='red', ls=':', alpha=0.5, label='50% 룰')
        ax.legend(fontsize=9)
        ax.grid(alpha=0.3)
    fig.tight_layout()
    fig.savefig(outpath, dpi=130)
    plt.close(fig)


def _set_doc_fonts(doc, name='맑은 고딕'):
    for sname in ['Normal', 'Heading 1', 'Heading 2', 'Title']:
        try:
            stl = doc.styles[sname]
            stl.font.name = name
            rpr = stl.element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn('w:eastAsia'), name)
        except Exception:
            pass


def build_docx(data, charts, out_path):
    doc = Document()
    _set_doc_fonts(doc)
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.0)

    title = doc.add_heading('AIMVP 포트폴리오 — 정정 요약 리포트', level=0)
    title.runs[0].font.color.rgb = HEAD1
    sub = doc.add_paragraph()
    sub.add_run(f"연속체인 정정 엔진 기준  ·  기준일 {data['date_label']}  ·  "
                f"가격 최종 {data['price_last'].date()}  ·  53개월 트랙레코드").italic = True

    a, n = data['적극형'], data['중립형']

    doc.add_heading('1. 핵심 결과 (정정 후)', level=1).runs[0].font.color.rgb = HEAD1
    doc.add_paragraph(
        '월간 리밸런싱 포트폴리오의 누적 성과를 연속 일별 체인(rebalance-at-close)으로 산출. '
        '구버전 엔진은 매 리밸 경계의 turn-of-month 1일 수익률(~52일)을 누락해 누적을 '
        '체계적으로 과소계상했으나, 본 리포트는 정정 엔진 기준이다.')
    tbl = doc.add_table(rows=3, cols=5)
    tbl.style = 'Light Grid Accent 1'
    hdr = ['Risk Profile', '실제(0%)', '50% 룰', '알파', 'Sharpe(0%→50%)']
    for j, h in enumerate(hdr):
        c = tbl.cell(0, j); c.text = h
        c.paragraphs[0].runs[0].font.bold = True
    rows = [
        ['적극형 (주식 89.6%)', f"+{a['itd_actual']:.2f}%", f"+{a['itd_swap']:.2f}%",
         f"+{a['alpha']:.2f}%p", f"{a['sharpe_actual']:.2f} → {a['sharpe_swap']:.2f}"],
        ['중립형 (주식 59.7%)', f"+{n['itd_actual']:.2f}%", f"+{n['itd_swap']:.2f}%",
         f"+{n['alpha']:.2f}%p", f"{n['sharpe_actual']:.2f} → {n['sharpe_swap']:.2f}"],
    ]
    for i, r in enumerate(rows, start=1):
        for j, v in enumerate(r):
            tbl.cell(i, j).text = v
    doc.add_paragraph(
        f"→ 중립이 적극의 약 {n['alpha']/a['alpha']*100:.0f}% 알파 (SPY/QQQ 비중 비례). "
        f"50% 스왑 룰은 두 프로파일 모두 (+)알파·Sharpe 개선.")

    doc.add_heading('2. 누적 성과 추이', level=1).runs[0].font.color.rgb = HEAD1
    for profile, key in [('적극형', '적극형'), ('중립형', '중립형')]:
        doc.add_heading(f'2.{1 if profile=="적극형" else 2} {profile}', level=2).runs[0].font.color.rgb = HEAD2
        doc.add_picture(charts[key], width=Cm(17))
        d = data[key]
        doc.add_paragraph(
            f"CAGR {d['cagr_actual']:+.2f}% · 변동성 {d['vol_actual']:.2f}% · "
            f"MaxDD {d['mdd_actual']:.2f}% · Sharpe {d['sharpe_actual']:.2f} (실제 0% 기준).")

    doc.add_heading('3. 스왑 비율 스윕 (0~100%)', level=1).runs[0].font.color.rgb = HEAD1
    doc.add_picture(charts['sweep'], width=Cm(17))
    doc.add_paragraph(
        'Bull 시그널 시 SPY→QQQ 스왑 비율을 0~100%로 변화시킨 누적/Sharpe 민감도. '
        '약관은 50% 룰을 명시하며, 100% 풀스왑은 2022~26 AI-rally 표본 편향 위험이 있어 보수적으로 본다.')

    doc.add_heading('4. Bull 시그널 적중률 (엔진 정정과 무관)', level=1).runs[0].font.color.rgb = HEAD1
    doc.add_paragraph(
        f"Bull 시그널 QQQ vs SPY 1개월 적중률: {a['hits']}/{a['n_bull']} "
        f"({a['hits']/a['n_bull']*100:.1f}%), 평균 초과 {a['avg_excess']:+.2f}%p. "
        "월별 독립 산출이라 누적 엔진 정정의 영향을 받지 않음 (적극·중립 동일). "
        "Wilcoxon p≈0.03으로 5% 유의 (기존 분석 유지).")

    doc.add_heading('5. 한계', level=1).runs[0].font.color.rgb = HEAD1
    for cav in [
        '표본 외 위험: 2022~2026은 AI-rally 편향 — 닷컴/금리쇼크형 SPY-우세 regime 미관측.',
        '검정력 한계: Bull n=21로 추가 트랙레코드 12~18개월 누적 후 재검증 권고.',
        '거래비용 미반영(그로스): QQQ-SPY 보수차 +11bps/년 + 양도세 + 스프레드 → net alpha 별도 측정 필요.',
    ]:
        doc.add_paragraph(cav, style='List Bullet')

    doc.add_paragraph()
    foot = doc.add_paragraph()
    foot.add_run(
        f"생성: generate_report.py · {data['gen_stamp']} · 데이터 출처 Yahoo Finance "
        "(auto-adjusted Total Return) + AIMVP 포트폴리오 비중추이.xlsx").italic = True

    doc.save(out_path)
    return out_path


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--date', default=None, help='기준일 YYYY-MM-DD (기본: 오늘)')
    args = ap.parse_args()
    end_str = args.date or datetime.now().strftime('%Y-%m-%d')

    print(f'[1/4] app 함수 로드...')
    g = load_app()
    print(f'[2/4] 수치 계산 (기준일 {end_str})...')
    data = compute_all(g, end_str)
    data['_g'] = g
    data['end_str'] = end_str
    price_label = str(data['price_last'].date())   # 실제 가격 최종일 (라벨/파일명 기준)
    data['date_label'] = price_label
    data['gen_stamp'] = datetime.now().strftime('%Y-%m-%d %H:%M')

    print(f'[3/4] 차트 생성...')
    charts = {
        '적극형': os.path.join(APP_DIR, 'report_적극형_cumulative.png'),
        '중립형': os.path.join(APP_DIR, 'report_중립형_cumulative.png'),
        'sweep': os.path.join(APP_DIR, 'report_swap_sweep.png'),
    }
    chart_cumulative(data, '적극형', charts['적극형'])
    chart_cumulative(data, '중립형', charts['중립형'])
    chart_swap_sweep(data, charts['sweep'])

    print(f'[4/4] docx 생성...')
    out_docx = os.path.join(APP_DIR, f'AIMVP_정정요약_{price_label}.docx')
    build_docx(data, charts, out_docx)

    a, n = data['적극형'], data['중립형']
    print('\n=== 생성 완료 ===')
    print(f"  적극형: 실제 +{a['itd_actual']:.2f}% / 50%룰 +{a['itd_swap']:.2f}% / 알파 +{a['alpha']:.2f}%p / Sharpe {a['sharpe_actual']:.2f}→{a['sharpe_swap']:.2f}")
    print(f"  중립형: 실제 +{n['itd_actual']:.2f}% / 50%룰 +{n['itd_swap']:.2f}% / 알파 +{n['alpha']:.2f}%p / Sharpe {n['sharpe_actual']:.2f}→{n['sharpe_swap']:.2f}")
    print(f"  charts: {os.path.basename(charts['적극형'])}, {os.path.basename(charts['중립형'])}, {os.path.basename(charts['sweep'])}")
    print(f"  docx:   {os.path.basename(out_docx)}")


if __name__ == '__main__':
    main()
